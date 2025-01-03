import json
import requests
import pickle
import os
import arxiv
# from config import Paper
from fuzzywuzzy import fuzz
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import logging
from download_pdf import get_pdf

# GetPDF is a variable that controls whether to spider PDF files.
GetPDF=True


def display_cit(cit):
    logging.info(
        "+++===================================================================================================+++"
    )
    logging.info(f"index: {cit['index']}")
    logging.info(f"title: {cit['title']}")
    logging.info(f"filename: {cit['filename']}")
    logging.info(f"info: {cit['info']}")
    logging.info(f"abstract: {cit['abstract']}")
    if cit['PDF'] == "":
        logging.info("no PDF resource.")
    else:
        logging.info(f"PDF: {cit['PDF']}")
    logging.info(f"paper_link: {cit['link']}")


# word format related functions
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """

    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        hs.font.underline = True
        hs.font.size = Pt(13)
        hs.font.name = 'Arial'
        del hs

    return "Hyperlink"

def get_locallink(paper, pdf_list):
    # paper_file=get_filename(paper.title)
    for pdf in pdf_list:
        ismatch=are_strings_almost_matching(paper.filename, pdf[:-4],threshold=90)
        if ismatch:
            return pdf
    return ""


def input_docx(cit, doc_pth, is_pdf, pdf_list=[]):
    logging.info("+======writing item======+")

    doc = Document(doc_pth)
    is_written = bool(len(doc.paragraphs))
    para = doc.add_paragraph()
    # set the space before the paragraph
    if is_written:
        para.paragraph_format.space_before = Pt(16)
    # set the first line indent
    para.paragraph_format.first_line_indent = Pt(0)
    # set the hanging indent
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    line1_text = cit['title']+'\n'
    if is_pdf:
        line0 = "[PDF downloaded]\n"
        run0 = para.add_run(line0)
        run0.font.name = "Arial"
        run0.font.size = Pt(12)
        run0.font.color.rgb = RGBColor(0, 200, 0) # green 
        line1_link = cit['filename'] + '.pdf'
    else:
        pdf_link=get_locallink(cit,pdf_list)
        if pdf_link:
            line0 = "[PDF downloaded]\n"
            run0 = para.add_run(line0)
            run0.font.name = "Arial"
            run0.font.size = Pt(12)
            run0.font.color.rgb = RGBColor(0, 200, 0) # green 
            line1_link = pdf_link
        else:
            line1_link = cit['link']

    add_hyperlink(para, line1_text,
                  line1_link)

    line2 = cit['info']+'\n'
    run2 = para.add_run(line2)
    run2.font.name = 'Arial'
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0, 102, 33)  # green

    line3 = cit['abstract']
    run3 = para.add_run(line3)
    run3.font.name = 'Arial'
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(34, 34, 34)  # black

    doc.save(doc_pth)

    logging.info("+======item done======+")

def are_strings_almost_matching(string1, string2, threshold=90):
    # 使用 fuzz.ratio() 比较字符串相似性
    similarity_ratio = fuzz.ratio(string1, string2)
    return similarity_ratio >= threshold


def get_filename(paper_title):
    words = paper_title.split()
    if len(words) <= 5:
        fn = paper_title
    else:
        fn = ' '.join(words[:5])
    fn = fn.replace(":", "")
    fn = fn.replace("?", "")
    return fn


def get_citation(dir_name, file_name):
    pth = f'./paper_list/{dir_name}/data/{file_name}'
    with open(pth, "rb") as file:
        paper = pickle.load(file)
    return paper


def list_data_in_directory(dir_name):
    # 使用 os.listdir 获取文件夹内所有文件和子文件夹的列表
    folder_path = f'./paper_list/{dir_name}/data/'
    files = os.listdir(folder_path)
    files.sort()
    return files

# def get_position(paper_list):
#     folder_path = f'./paper_list/'
#     paper_processed = os.listdir(folder_path)
#     if not paper_processed:
#         return 0
#     c = 0
#     for ind, title in enumerate(paper_list):
#         dir_name = get_filename(title)
#         docx_path = folder_path+f'{dir_name}/{dir_name}.docx'
#         if os.path.exists(docx_path):
#             c = ind
#     return c


def docx_worker(paper_title):
    print(f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***")
    print(f"Start writing the docx document of the paper: [{paper_title}]")
    print(f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***\n")

    dir_name = get_filename(paper_title)
    if os.path.exists(f'./paper_list/{dir_name}/citation_info.json'):
        with open(f'./paper_list/{dir_name}/citation_info.json', 'r') as file:
            cit_list = json.load(file)
    else:
        cit_list = []

    # files = list_data_in_directory(dir_name)
    isPDF = 0

    logging.info("\n\n\n")
    logging.info(f"\n***+++++++++++++++++++++++++++++writing the docx of Paper: [{dir_name}]+++++++++++++++++++++++++++++***\n")
    if not cit_list:
        logging.info(f"Paper: [{dir_name}] has no citation")
        print(f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***")
        print(f"Paper: [{dir_name}] has no citation")
        print(f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***")
        print()
        return

    if GetPDF:
        doc_pth = f'./paper_list/{dir_name}/{dir_name}.docx'
    else:
        doc_pth = f'./paper_list/{dir_name}/(final) {dir_name}.docx'

    doc = Document()
    doc.save(doc_pth)

    if not GetPDF:
        dir_path=os.path.dirname(doc_pth)
        # 记录已经下载的pdf文件，用于修改相对地址
        pdf_files = [file for file in os.listdir(dir_path) if file.endswith('.pdf')]

    for cit in cit_list:

        display_cit(cit)
        pdf_pth = f"./paper_list/{dir_name}/{cit['filename']}.pdf"

        if GetPDF:
            isPDF = get_pdf(cit, pdf_pth)
            input_docx(cit, doc_pth, isPDF)
        else:
            input_docx(cit, doc_pth, False, pdf_list=pdf_files)
        logging.info("+++===================================================================================================+++\n")

    print(f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***")
    print(f"The docx document of the paper: [{paper_title}] has been written successfully.")
    print(f"***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***\n")


def docx_generator(paper_ls):   
    print()
    print(f"The {str(len(paper_ls))} docx documents to be written:")
    print(paper_ls)
    print("+++===================================================================================================+++")
    print()

    
    logging.info("\n\n\n")
    logging.info(f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####")
    logging.info(f"The following is a new process")
    logging.info(f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####")
    logging.info("\n\n\n")

    for paper in paper_ls:
        docx_worker(paper) 
    print("All docx documents have been written successfully.")

    logging.info("\n\n\n")
    logging.info(f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####")
    logging.info(f"All docx documents have been written successfully.")
    logging.info(f"#####***++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++***#####")
    logging.info("\n\n\n")

def get_papers():
    dir_ls=os.listdir("./paper_list")
    paper_ls = sorted([d for d in dir_ls if not (d.endswith('.log') or d.startswith('.'))])
    return paper_ls

if __name__ == "__main__":
    if os.path.exists("./paper_list"):
        paper_list=get_papers()
        # print(paper_list)
        log_filename = f'./paper_list/docx.log'
        logging.basicConfig(filename=log_filename, level=logging.INFO, format='%(message)s')
        docx_generator(paper_list)
    else:
        print("Please use CitationSpider to get citation data in advance")
