from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
import json


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    # 清除或者覆盖默认的 Hyperlink 样式
    # 使得后续的手动颜色设置能生效
    r_pr = OxmlElement("w:rPr")
    # 强制指定纯色
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1A0DAB")  # 1A0DAB等价于RGB(26, 13, 171)
    r_pr.append(color_el)

    # 字体大小
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")  # 10磅，对应pt(10)
    r_pr.append(sz)

    # 字体名称
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_pr.append(r_fonts)

    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# 加载 JSON 数据
json_file_path = "Zhouchen_Lin.json"
with open(json_file_path, "r", encoding="utf-8") as file:
    data = json.load(file)

doc = Document()

doc.add_heading("2024年引用摘录分工", level=1)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

# ---- 将表格线改为虚线 ----
tbl = table._tbl
tblPr = tbl.tblPr
tblBorders = tblPr.first_child_found_in("w:tblBorders")

if tblBorders is None:
    tblBorders = OxmlElement("w:tblBorders")
    tblPr.append(tblBorders)

for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
    border_el = tblBorders.find(qn("w:%s" % border_name))
    if border_el is None:
        border_el = OxmlElement("w:%s" % border_name)
        tblBorders.append(border_el)
    # 设为虚线 (dotted)，线宽可根据需要进行调整
    border_el.set(qn("w:val"), "dotted")
    border_el.set(qn("w:sz"), "4")
    border_el.set(qn("w:space"), "0")
    border_el.set(qn("w:color"), "auto")
# ---- 虚线设置结束 ----


# 手动指定列宽
table.autofit = False
table.columns[0].width = Inches(5)
table.columns[1].width = Inches(0.7)
table.columns[2].width = Inches(0.7)

# 设置表头并居中
headers = ["论文", "引用数", "负责人"]
header_row = table.rows[0]
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = 1  # 居中对齐

# 填充表格内容
for entry in data:
    row_cells = table.add_row().cells

    # 第1列（信息）
    info_cell = row_cells[0]
    info_paragraph = info_cell.paragraphs[0]
    info_paragraph.paragraph_format.line_spacing = Pt(12)  # 调小行距
    info_paragraph.paragraph_format.space_before = Pt(0)
    info_paragraph.paragraph_format.space_after = Pt(0)

    title = entry.get("title", "N/A")
    authors = entry.get("authors", "N/A")
    publication = entry.get("publication", "N/A")
    link = entry.get("link", "#")

    # 添加标题为超链接
    add_hyperlink(info_paragraph, title, link)

    # 插入作者
    info_paragraph.add_run("\n" + authors).font.size = Pt(9)
    # 插入发表信息
    info_paragraph.add_run("\n" + publication).font.size = Pt(9)

    # 第2列：引用次数（居中）
    cited_by = str(entry.get("cite_num_within_time", 0))
    row_cells[1].text = cited_by
    for p in row_cells[1].paragraphs:
        p.alignment = 1  # 居中对齐
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)

    # 第3列：负责人（居中）
    row_cells[2].text = ""
    for p in row_cells[2].paragraphs:
        p.alignment = 1  # 居中对齐
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)

output_file_path = "2024年引用摘录分工_zlin.docx"
doc.save(output_file_path)

print(f"文档已保存到 {output_file_path}")
